from typing import List
from docx import Document


def read_docx(path: str) -> List[str]:
    """
    Reads the docx file and returns its paragraphs as a list of text.

    Inputs:
    ----------
        path (str):
            The path to the .docx file.

    Returns:
    -------
        List[str]
            A list of paragraphs from the document.
    """

    doc = Document(path)

    temp = []
    for para in doc.paragraphs:
        temp.append(para.text)
    paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
    return paragraphs

